# -*- coding: utf-8 -*-
"""
Created on Fri Aug 04 16:16:14 2023

@author: Long_S
"""

import pandas as pd
import sqlalchemy as sa
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document

from matplotlib import rcParams
from matplotlib import pyplot as plt

def data_table(group, digit, yr, abbreviated, export):
    try:
        engine = sa.create_engine(
        "mssql+pyodbc://OPTSRV2/IPS2DB_SAS?driver=ODBC+Driver+17+for+SQL+Server")
        driver_test = pd.read_sql("""SELECT TOP 1 * FROM dbo.ReportBuilder_Preliminary""", engine)
    except:
        engine = sa.create_engine(
            "mssql+pyodbc://OPTSRV2/IPS2DB_SAS?driver=ODBC Driver 18 for SQL Server&ssl_verify_cert=false&TrustServerCertificate=yes")
        
    ## Get industry groups
    industry_groups = pd.read_sql("""SELECT IndustryID, IndustryGroupID FROM dbo.IndustryGroupSets""", engine)
    industry_groups = industry_groups[~industry_groups['IndustryGroupID'].isin(['BEANIPA', 'Sector63'])]
    
    doc = Document()
    
    p1 = doc.add_paragraph("Tables were generated with the following parameters:")
    p1.add_run().add_break()
    p1.add_run(f"Industry Group = {group}, Industry Digit = {digit}, Year = {yr}, Abbreviated variable names = {abbreviated}, export = {export}]")

    p2 = doc.add_paragraph("This can be replicated with the command:")
    p2.add_run().add_break()
    p2.add_run(f"data_in_year({group}, {digit}, {yr}, {abbreviated}, {export})") 
    
    doc.add_page_break()    
    
    ## Get industries from specified group
    industries = tuple(industry_groups[industry_groups['IndustryGroupID'] == group]['IndustryID'])
    
    df = pd.read_sql(f"""
                     SELECT Industry, IndustryTitle, DataSeries, DataSeriesTitle, IndustryDigit, Year, Unit, Value 
                     FROM dbo.ReportBuilder_Current
                     WHERE IndustryID IN {industries} AND IndustryDigit = '{digit}' AND PublicAccess ='True' AND (Year = {yr} OR Year = {yr-1})""", engine)
    
    dataframes = []

    for title in df['IndustryTitle'].unique():
        dfv2 = df.loc[df['IndustryTitle'] == title]
        
        naics = dfv2['Industry'].unique()[0]

        output_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output index') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)

        hours_df = dfv2[(dfv2['DataSeriesTitle'] == 'Hours index') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)

        oph_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output per hour') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True) 

        val_df = dfv2[(dfv2['DataSeriesTitle'] == 'Value of production') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)  

        ipd_df = dfv2[(dfv2['DataSeriesTitle'] == 'Implicit price deflator') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True) 

        ulc_df = dfv2[(dfv2['DataSeriesTitle'] == 'Unit labor cost index') & (dfv2['Year'] >= yr-1) & (dfv2['Year'] <= yr)] \
                    .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True) 
        
        output_df['Year percent change'] = output_df['Value'].pct_change() * 100
        hours_df['Year percent change'] = hours_df['Value'].pct_change() * 100
        oph_df['Year percent change'] = oph_df['Value'].pct_change() * 100
        val_df['Year percent change'] = val_df['Value'].pct_change() * 100
        ipd_df['Year percent change'] = ipd_df['Value'].pct_change() * 100
        ulc_df['Year percent change'] = ulc_df['Value'].pct_change() * 100
                       
        dataframes.append((naics, title, output_df, hours_df, oph_df, val_df, ipd_df, ulc_df))
    
    table_title = doc.add_paragraph(f"Table of {group} {digit} Industries Percent Change Metrics in {yr}")
    table_title.add_run().add_break()
    
    table = doc.add_table(rows=len(dataframes)+1, cols=len(dataframes[0]))
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = 'NAICS'
    table.rows[0].cells[1].text = 'Industry'
    if not abbreviated:
        for i in range(len(dataframes[0])-2):
            table.rows[0].cells[i+2].text = str(dataframes[0][i+2]['DataSeriesTitle'].unique()[0]).replace(' index', '')
    else:
        for i in range(len(dataframes[0])-2):
            table.rows[0].cells[i+2].text = str(dataframes[0][i+2]['DataSeries'].unique()[0])
    
    for i, df in enumerate(dataframes):
        table.rows[i+1].cells[0].text = df[0]
        table.rows[i+1].cells[1].text = df[1]
        for k in range(len(dataframes[0])-2):
            table.rows[i+1].cells[k+2].text = str(round(df[k+2]['Year percent change'][1], 2))

    filepath = export[1:-1]
    doc.save(filepath + '\\data_table.docx')