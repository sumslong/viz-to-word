# -*- coding: utf-8 -*-
"""
Created on Tue Jul 18 16:16:14 2023

@author: Long_S
"""

import pandas as pd
import sqlalchemy as sa
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document

from matplotlib import rcParams
from matplotlib import pyplot as plt


def hours_output_prod_chart(group, periods=False, start_yr=2009, end_yr=2022, sort='extreme oph change first', digit='3-Digit', limit=5, export='labor_productivity_charts.docx'):
    
    try:
        engine = sa.create_engine(
        "mssql+pyodbc://OPTSRV2/IPS2DB_SAS?driver=ODBC+Driver+17+for+SQL+Server")
        driver_test = pd.read_sql("""SELECT TOP 1 * FROM dbo.ReportBuilder_Preliminary""", engine)
    except:
        engine = sa.create_engine(
            "mssql+pyodbc://OPTSRV2/IPS2DB_SAS?driver=ODBC Driver 18 for SQL Server&ssl_verify_cert=false&TrustServerCertificate=yes")
        
    ## Get industry groups
    industry_groups = pd.read_sql("""SELECT IndustryID, IndustryGroupID FROM dbo.IndustryGroupSets""", engine)
    industry_groups = industry_groups[~industry_groups['IndustryGroupID'].isin(['BEANIPA', 'Sector63'])]
    
    doc = Document()
    
    p1 = doc.add_paragraph("Plots were generated with the following parameters:")
    p1.add_run().add_break()
    p1.add_run(f"Industry Group = {group}, Periods = {periods}, Start year = {start_yr}, End year = {end_yr}, Sort by = {sort}, Industry Digit = {digit}, Number of graphs = {limit}")
    
    p2 = doc.add_paragraph("This can be replicated with the command:")
    p2.add_run().add_break()
    p2.add_run(f"hours_output_prod_chart('{group}', '{periods}', {start_yr}, {end_yr}, '{sort}', '{digit}', {limit}. {export})") 
        
    if periods:
        doc.add_paragraph("The data was grouped into periods and annual percent change was computed within a period for each measure.")
        
    if not periods:
        doc.add_paragraph("The year to year percent change was computed for each measure.")
        
    if sort == 'extreme oph change first':
        doc.add_paragraph("The difference in output per hour percent change between the most recent data point and the data point prior was computed. The charts were then sorted to display in descending order.")
        
    if sort == 'highest relative oph change first': 
        doc.add_paragraph("The difference in output per hour percent change between the most recent data point and the average output per hour percent change up to the most recent data point was computed. The charts were then sorted to display in descending order.")
        
    if sort == 'highest interest score first':
        p3 = doc.add_paragraph("An interest score was calculated with the following equation:")
        p3.add_run().add_break()
        p3.add_run("(|.2x| + |.2y| + |.6z|) * b")
        p3.add_run().add_break()
        p3.add_run("Where x is the difference between the percent change in the output index for the 2 most recent datapoints recorded,").add_break()
        p3.add_run("y is the difference between the percent change in the hours index for the 2 most recent datapoints recorded,").add_break()
        p3.add_run("z is the difference between the percent change in the output per hour index for the 2 most recent datapoints recorded,").add_break()
        p3.add_run("b is the number of employees in the most recent datapoint recorded")
        doc.add_paragraph("The charts were then sorted to display in descending order based on calculated interest score.")        
        
    ## Get industries from specified group
    industries = tuple(industry_groups[industry_groups['IndustryGroupID'] == group]['IndustryID'])
    
    df = pd.read_sql(f"""
                     SELECT IndustryTitle, DataSeriesTitle, IndustryDigit, Year, Unit, Value 
                     FROM dbo.ReportBuilder_Current
                     WHERE IndustryID IN {industries} AND IndustryDigit = '{digit}' AND PublicAccess ='True'""", engine)
    
    ## Some industries do not have sufficient information to create plot. Find these for later exclusion in loop.
    expected = {'Output index', 'Hours index', 'Output per hour'}
    missing_data_series = df.groupby('IndustryTitle')['DataSeriesTitle'] \
                            .apply(lambda x: expected - set(x)) \
                            .loc[lambda x: x.apply(lambda s: bool(s))] \
                            .reset_index()
    
    ## List for storing df
    dataframes = []

    ## Loop for plot creation: one plot for each industry
    for title in df['IndustryTitle'].unique():
        ## Only plot ones that have sufficient information
        if title not in missing_data_series['IndustryTitle'].values:
            if sort == 'extreme oph change first' or sort == 'highest relative oph change first':
                dfv2 = df.loc[(df['IndustryTitle'] == title) & ((df['DataSeriesTitle'] == 'Output index') | \
                       (df['DataSeriesTitle'] == 'Hours index') | (df['DataSeriesTitle'] == 'Output per hour'))] 
                
            if sort == 'highest interest score first':
                dfv2 = df.loc[(df['IndustryTitle'] == title) & ((df['DataSeriesTitle'] == 'Output index') | \
                           (df['DataSeriesTitle'] == 'Hours index') | (df['DataSeriesTitle'] == 'Output per hour') \
                            | (df['DataSeriesTitle'] == 'All persons'))] 
                       
            ## If plot is yearly
            if not periods:     
                ## Only include years specified
                
                output_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output index') & (dfv2['Year'] >= start_yr-1) & (dfv2['Year'] <= end_yr)] \
                            .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                
                hours_df = dfv2[(dfv2['DataSeriesTitle'] == 'Hours index') & (dfv2['Year'] >= start_yr-1) & (dfv2['Year'] <= end_yr)] \
                            .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                
                oph_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output per hour') & (dfv2['Year'] >= start_yr-1) & (dfv2['Year'] <= end_yr)] \
                            .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                
                emp_df = dfv2[(dfv2['DataSeriesTitle'] == 'All persons') & (dfv2['Year'] >= start_yr-1) & (dfv2['Year'] <= end_yr)] \
                            .sort_values('Year', ascending=True).drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)  
               
                ## Compute percent change
                output_df['Year percent change'] = output_df['Value'].pct_change() * 100
                hours_df['Year percent change'] = hours_df['Value'].pct_change() * 100
                oph_df['Year percent change'] = oph_df['Value'].pct_change() * 100
                       
                if sort == 'extreme oph change first':
                    rank = oph_df['Year percent change'].iloc[-2] - oph_df['Year percent change'].iloc[-1]
                    dataframes.append([title, output_df, hours_df, oph_df, abs(rank)])

                if sort == 'highest relative oph change first':
                    rank = oph_df['Year percent change'].iloc[:-2].mean() - oph_df['Year percent change'].iloc[-1]
                    dataframes.append([title, output_df, hours_df, oph_df, abs(rank)])
                       
                if sort == 'highest interest score first':
                    rank = (abs(.6*(oph_df['Year percent change'].iloc[-2] - oph_df['Year percent change'].iloc[-1])) \
                           + abs(.2*(hours_df['Year percent change'].iloc[-2] - hours_df['Year percent change'].iloc[-1])) \
                           + abs(.2*(output_df['Year percent change'].iloc[-2] - output_df['Year percent change'].iloc[-1]))) \
                           * emp_df['Value'].iloc[-1]
                    dataframes.append([title, output_df, hours_df, oph_df, rank])
                    
            if periods:
                if 1987 in dfv2['Year'].values:
                    ## Periods for plotting
                    bins = [
                        (1987, 1990),
                        (1990, 2000),
                        (2000, 2007),
                        (2007, 2019),
                        (2019, dfv2['Year'].max())
                        ]

                    results_output = []
                    results_hours = []
                    results_oph = []

                    ## Need all years
                    output_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output index')].sort_values('Year', ascending=True) \
                                .drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                    hours_df = dfv2[(dfv2['DataSeriesTitle'] == 'Hours index')].sort_values('Year', ascending=True) \
                                .drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                    oph_df = dfv2[(dfv2['DataSeriesTitle'] == 'Output per hour')].sort_values('Year', ascending=True) \
                                .drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                    if sort == 'highest interest score first':
                        emp_df = dfv2[(dfv2['DataSeriesTitle'] == 'All persons')].sort_values('Year', ascending=True) \
                                  .drop_duplicates(subset=['Year', 'Value']).reset_index(drop=True)
                        start, end = bins[4]
                        avg_emp = emp_df[(emp_df['Year'] >= start) & (emp_df['Year'] <= end)]['Year'].mean()
                    
                    if 1987 in output_df['Year'].values and 1987 in hours_df['Year'].values and 1987 in oph_df['Year'].values:
                        ## Compute annual percent change
                        for start_year, end_year in bins:
                            filtered_output = output_df[(output_df['Year'] >= start_year) & (output_df['Year'] <= end_year)].reset_index(drop=True)
                            filtered_hours = hours_df[(hours_df['Year'] >= start_year) & (hours_df['Year'] <= end_year)].reset_index(drop=True)
                            filtered_oph = oph_df[(oph_df['Year'] >= start_year) & (oph_df['Year'] <= end_year)].reset_index(drop=True)

                            result_output = (f'{start_year}-{end_year}', 
                                            ((filtered_output['Value'].iloc[0] / filtered_output['Value'].iloc[-1]) ** (1 / len(filtered_output)) - 1) * 100)
                            result_hours = (f'{start_year}-{end_year}', 
                                            ((filtered_hours['Value'].iloc[0] / filtered_hours['Value'].iloc[-1]) ** (1 / len(filtered_hours)) - 1) * 100)
                            result_oph = (f'{start_year}-{end_year}', 
                                            ((filtered_oph['Value'].iloc[0] / filtered_oph['Value'].iloc[-1]) ** (1 / len(filtered_oph)) - 1) * 100)

                            results_output.append(result_output)
                            results_hours.append(result_hours)
                            results_oph.append(result_oph)
                        
                        output_org = output_df.tail(5).copy()
                        hours_org = hours_df.tail(5).copy()
                        
                        output_df = pd.DataFrame(results_output, columns=['Period', 'Annual percent change'])
                        hours_df = pd.DataFrame(results_hours, columns=['Period', 'Annual percent change'])
                        oph_df = pd.DataFrame(results_oph, columns=['Period', 'Annual percent change'])

                        if sort == 'extreme oph change first':
                            rank = oph_df['Annual percent change'].iloc[-2] - oph_df['Annual percent change'].iloc[-1]
                            dataframes.append([title, output_df, hours_df, oph_df, abs(rank)])

                        if sort == 'highest relative oph change first':
                            rank = oph_df['Annual percent change'].iloc[:-2].mean() - oph_df['Annual percent change'].iloc[-1]
                            dataframes.append([title, output_df, hours_df, oph_df, abs(rank)])

                        if sort == 'highest interest score first':
                            rank = (abs(.6*(oph_df['Annual percent change'].iloc[-2] - oph_df['Annual percent change'].iloc[-1])) \
                                    + abs(.2*(hours_df['Annual percent change'].iloc[-2] - hours_df['Annual percent change'].iloc[-1])) \
                                    + abs(.2*(output_df['Annual percent change'].iloc[-2] - output_df['Annual percent change'].iloc[-1]))) \
                                    * avg_emp
                            dataframes.append([title, output_df, hours_df, oph_df, rank])
                           
    if not periods:
        sorted_list = sorted(dataframes, key=lambda x: abs(x[4]), reverse=True)
        for df in sorted_list[:limit]:
            doc.add_page_break()
            x = np.arange(len(df[1]['Year'].sort_values(ascending=True))-1)
            bar_width = 0.35
            plt.bar(x, df[1]['Year percent change'][1:], width=bar_width, color='darkblue', label='Output index')
            plt.bar(x + bar_width, df[2]['Year percent change'][1:], width=bar_width, color='lightblue', label='Hours index')

            ## Output per hour points and line
            plt.plot(x + bar_width / 2, df[3]['Year percent change'][1:], color='red')
            plt.scatter(x + bar_width / 2, df[3]['Year percent change'][1:], color='black', marker='o', label='Output per hour', zorder=10)

            ## Year position
            plt.xticks(x + bar_width / 2, df[1]['Year'][1:], fontsize=8)
            plt.title(f"{df[0]}")
            plt.xlabel('Year')
            plt.ylabel('Year percent change')
            plt.legend(bbox_to_anchor=(.5, -.15), loc='upper center', ncol=3)
            plt.plot()

            ## Save to word
            plot_filename = f'{df[0]}.png'
            plt.savefig(plot_filename, bbox_inches='tight')
            doc.add_picture(plot_filename)
            
                        
            table = doc.add_table(rows=len(output_df), cols=4)
            table.style = 'Table Grid'
            
            for i in range(3):
                table.rows[0].cells[i+1].text = str(df[i+1]['DataSeriesTitle'].unique()[0])
            for i in range(len(output_df)-1):
                table.rows[i+1].cells[0].text = str(list(df[1]['Year'][1:])[i])
                table.rows[i+1].cells[1].text = str(round(list(df[1]['Year percent change'][1:])[i], 1))
                table.rows[i+1].cells[2].text = str(round(list(df[2]['Year percent change'][1:])[i], 1))
                table.rows[i+1].cells[3].text = str(round(list(df[3]['Year percent change'][1:])[i], 1))
                
            os.remove(plot_filename)

            plt.clf()

    if periods:
        sorted_list = sorted(dataframes, key=lambda x: abs(x[4]), reverse=True)
        for df in sorted_list[:limit]:
            doc.add_page_break()
            x = np.arange(len(df[1]['Period'].sort_values(ascending=True)))
            bar_width = 0.35
            plt.bar(x, df[1]['Annual percent change'], width=bar_width, color='darkblue', label='Output index')
            plt.bar(x + bar_width, df[2]['Annual percent change'], width=bar_width, color='lightblue', label='Hours index')

            ## Output per hour points and line
            plt.plot(x + bar_width / 2, df[3]['Annual percent change'], color='red')
            plt.scatter(x + bar_width / 2, df[3]['Annual percent change'], color='black', marker='o', label='Output per hour', zorder=10)

            ## Year position
            plt.xticks(x + bar_width / 2, df[1]['Period'], fontsize=8)
            plt.title(f"{df[0]}")
            plt.xlabel('Period')
            plt.ylabel('Annual percent change')
            plt.legend(bbox_to_anchor=(.5, -.15), loc='upper center', ncol=3)
            plt.plot()

            ## Save to word
            plot_filename = f'{df[0]}.png'
            plt.savefig(plot_filename, bbox_inches='tight')
            doc.add_picture(plot_filename)

            left_margin = (doc.sections[0].page_width - doc.inline_shapes[-1].width) / 2
            top_margin = (doc.sections[0].page_height - doc.inline_shapes[-1].height) / 2

            doc.inline_shapes[-1].left = left_margin
            doc.inline_shapes[-1].top = top_margin
            
            table = doc.add_table(rows=len(output_df)+1, cols=4)
            table.style = 'Table Grid'
            
            table.rows[0].cells[1].text = "Output index"
            table.rows[0].cells[2].text = "Hours index"
            table.rows[0].cells[3].text = "Output per hour" 

            for i in range(len(output_df)):
                table.rows[i+1].cells[0].text = str(list(df[1]['Period'])[i])
                table.rows[i+1].cells[1].text = str(round(list(df[1]['Annual percent change'])[i], 1))
                table.rows[i+1].cells[2].text = str(round(list(df[2]['Annual percent change'])[i], 1))
                table.rows[i+1].cells[3].text = str(round(list(df[3]['Annual percent change'])[i], 1))

            os.remove(plot_filename)

            plt.clf()
            
    filepath = export[1:-1]
    doc.save(filepath + '\\lp_graphs.docx')
